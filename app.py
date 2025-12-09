from collections import Counter  # (currently unused, but fine to keep)
import pandas as pd
import streamlit as st
import sqlite3
from io import BytesIO
from datetime import datetime, date, timedelta
import random
import calendar

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from agent.symptom_extractor import extract_symptoms
from agent.disease_engine import match_disease, explain_symptoms
from agent.disease_matcher import rank_diseases, DISEASE_KB
from medical_rules import analyze_vitals
from agent.doctor_engine import suggest_specialities, rank_doctors
from agent.followup_engine import estimate_risk_from_symptoms, get_followup_plan
from agent.treatment_engine import get_care_tips


# =====================================================
# SMALL HELPERS
# =====================================================

def disease_name(d):
    """Safely get just the disease name from either a dict or a string."""
    if isinstance(d, dict) and "disease" in d:
        return d["disease"]
    return str(d)


# =====================================================
# DATABASE SETUP
# =====================================================

def init_db():
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()

    # --- Patient health records ---
    c.execute("""
        CREATE TABLE IF NOT EXISTS patients (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            age INTEGER,
            symptoms TEXT,
            bp TEXT,
            heart_rate TEXT,
            temperature TEXT,
            risk_score TEXT,
            possible_diseases TEXT
        )
    """)

    # Medication adherence history (optional, not heavily used right now)
    c.execute("""
        CREATE TABLE IF NOT EXISTS med_adherence (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_name TEXT,
            medication_name TEXT,
            date TEXT,
            taken INTEGER   -- 1 = taken, 0 = missed
        )
    """)

    # Daily wellness logs
    c.execute("""
        CREATE TABLE IF NOT EXISTS wellness_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_name TEXT,
            steps INTEGER,
            sleep_hours REAL,
            water_intake INTEGER,
            mood TEXT,
            date TEXT
        )
    """)

    # Wellness goals (for wellness tracking)
    c.execute("""
        CREATE TABLE IF NOT EXISTS wellness_goals (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_name TEXT,
            steps_goal INTEGER,
            sleep_goal REAL,
            water_goal INTEGER
        )
    """)

    conn.commit()
    conn.close()

# =====================================================
# VALIDATION HELPERS
# =====================================================

def validate_vitals_inputs(bp: str, heart_rate: str, temperature: str):
    """Return (ok: bool, errors: list[str]) for vitals."""
    errors = []

    # BP: simple format check: "120/80"
    if "/" not in bp:
        errors.append("Blood pressure should be in the format 120/80.")
    else:
        parts = bp.split("/")
        if len(parts) != 2:
            errors.append("Blood pressure should have two numbers like 120/80.")
        else:
            try:
                sys = int(parts[0].strip())
                dia = int(parts[1].strip())
                if not (70 <= sys <= 250 and 40 <= dia <= 150):
                    errors.append(
                        "Blood pressure values look unusual; please re-check them."
                    )
            except ValueError:
                errors.append("Blood pressure values should be numbers (e.g., 120/80).")

    # Heart rate numeric + reasonable
    try:
        hr_val = int(heart_rate)
        if not (30 <= hr_val <= 220):
            errors.append("Heart rate should normally be between 30 and 220 bpm.")
    except ValueError:
        errors.append("Heart rate should be a number (bpm).")

    # Temperature numeric
    try:
        float(temperature)
    except ValueError:
        errors.append("Temperature should be a number in °C (e.g., 37.5).")

    return (len(errors) == 0, errors)


def validate_medication_inputs(med_patient, med_name, med_dosage,
                               med_times, med_start, med_end):
    """Return (ok: bool, errors: list[str]) for medication scheduling."""
    errors = []
    if not med_patient.strip():
        errors.append("Patient name is required for medication schedule.")
    if not med_name.strip():
        errors.append("Medication name cannot be empty.")
    if not med_dosage.strip():
        errors.append("Please enter dosage (e.g., 500 mg).")
    if not med_times:
        errors.append("Select at least one time of day for this medication.")
    if med_end < med_start:
        errors.append("End date cannot be before start date.")

    return (len(errors) == 0, errors)



def init_medication_db():
    """Medication schedule (for adherence tracking)."""
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()

    # Just create if not exists – DO NOT DROP EVERY TIME
    c.execute("""
        CREATE TABLE IF NOT EXISTS medications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_name TEXT,
            medication_name TEXT,
            dosage TEXT,
            frequency TEXT,      -- Once a day / Twice a day / Thrice a day
            time_of_day TEXT,    -- "Morning, Night" etc.
            start_date TEXT,
            end_date TEXT
        )
    """)
    conn.commit()
    conn.close()


def init_medication_log_db():
    """Logs for each taken dose."""
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS medication_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            medication_id INTEGER,
            log_date TEXT,       -- YYYY-MM-DD
            time_of_day TEXT,    -- Morning / Night etc.
            status TEXT          -- 'taken'
        )
    """)
    conn.commit()
    conn.close()


def reset_appointment_table():
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("DROP TABLE IF EXISTS appointments")
    conn.commit()
    conn.close()


def init_appointments_db():
    """Appointments table (per-doctor schedule)."""
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS appointments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            doctor TEXT,
            speciality TEXT,
            patient TEXT,
            date TEXT,
            time TEXT
        )
        """
    )
    conn.commit()
    conn.close()


def row_to_patient(row):
    keys = [
        "name",
        "age",
        "symptoms",
        "bp",
        "heart_rate",
        "temperature",
        "risk_score",
        "possible_diseases",
    ]
    return dict(zip(keys, row))


def fetch_latest_patient(name: str):
    """Get the latest analysis for a given patient name (case-insensitive)."""
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute(
        """
        SELECT symptoms, risk_score
        FROM patients
        WHERE LOWER(name) = LOWER(?)
        ORDER BY id DESC
        LIMIT 1
        """,
        (name.strip(),),
    )
    result = c.fetchone()
    conn.close()
    return result


# ---- actually run DB initialisation (after all defs) ----
init_db()
init_medication_db()
init_medication_log_db()
reset_appointment_table()
init_appointments_db()


# =====================================================
# REPORT GENERATORS (DOCX / PDF)
# =====================================================

def create_docx_for_patient(patient):
    doc = Document()
    doc.add_heading("Healthcare Monitoring Report", level=1)

    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    doc.add_paragraph(f"Name: {patient['name']}")
    doc.add_paragraph(f"Age: {patient['age']}")
    doc.add_paragraph(f"Symptoms: {patient['symptoms']}")
    doc.add_paragraph(f"Blood Pressure: {patient['bp']}")
    doc.add_paragraph(f"Heart Rate: {patient['heart_rate']} bpm")
    doc.add_paragraph(f"Temperature: {patient['temperature']} °C")
    doc.add_paragraph(f"Risk Level: {patient['risk_score']}")
    doc.add_paragraph(f"Possible Diseases: {patient['possible_diseases']}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_docx_for_all(patients):
    doc = Document()
    doc.add_heading("All Patient Health Records", level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    for idx, p in enumerate(patients, start=1):
        doc.add_heading(f"Patient {idx}: {p['name']}", level=2)
        doc.add_paragraph(f"Age: {p['age']}")
        doc.add_paragraph(f"Symptoms: {p['symptoms']}")
        doc.add_paragraph(f"Blood Pressure: {p['bp']}")
        doc.add_paragraph(f"Heart Rate: {p['heart_rate']} bpm")
        doc.add_paragraph(f"Temperature: {p['temperature']} °C")
        doc.add_paragraph(f"Risk Level: {p['risk_score']}")
        doc.add_paragraph(f"Possible Diseases: {p['possible_diseases']}")
        doc.add_paragraph("-" * 40)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_pdf_for_patient(patient):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    y = height - 50
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, y, "Healthcare Monitoring Report")
    y -= 40

    c.setFont("Helvetica", 11)
    lines = [
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
        f"Name: {patient['name']}",
        f"Age: {patient['age']}",
        f"Symptoms: {patient['symptoms']}",
        f"Blood Pressure: {patient['bp']}",
        f"Heart Rate: {patient['heart_rate']} bpm",
        f"Temperature: {patient['temperature']} °C",
        f"Risk Level: {patient['risk_score']}",
        f"Possible Diseases: {patient['possible_diseases']}",
    ]

    for line in lines:
        c.drawString(50, y, line)
        y -= 18

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


def create_pdf_for_all(patients):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    for idx, patient in enumerate(patients, start=1):
        y = height - 50
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, f"Patient {idx}: {patient['name']}")
        y -= 30

        c.setFont("Helvetica", 11)
        lines = [
            f"Age: {patient['age']}",
            f"Symptoms: {patient['symptoms']}",
            f"Blood Pressure: {patient['bp']}",
            f"Heart Rate: {patient['heart_rate']} bpm",
            f"Temperature: {patient['temperature']} °C",
            f"Risk Level: {patient['risk_score']}",
            f"Possible Diseases: {patient['possible_diseases']}",
        ]

        for line in lines:
            c.drawString(50, y, line)
            y -= 18

        c.showPage()

    c.save()
    buffer.seek(0)
    return buffer

# =====================================================
# DASHBOARD HELPERS (VITALS + WELLNESS)
# =====================================================

def get_latest_vitals_for_dashboard(patient_name: str):
    """Get last saved vitals + risk for dashboard."""
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT bp, heart_rate, temperature, risk_score, symptoms
        FROM patients
        WHERE LOWER(name) = LOWER(?)
        ORDER BY id DESC
        LIMIT 1
    """, (patient_name.strip(),))
    row = c.fetchone()
    conn.close()

    if not row:
        return None

    return {
        "bp": row[0],
        "heart_rate": row[1],
        "temperature": row[2],
        "risk": row[3],
        "symptoms": row[4],
    }


def get_wellness_stats_for_dashboard(patient_name: str):
    """
    Returns (stats_dict, df) for wellness dashboard.

    stats_dict keys:
      - avg_steps, avg_sleep, avg_water
      - streak_days
      - goal_completion: {'steps': .., 'sleep': .., 'water': ..}
      - avg_goal_completion
    df: pandas DataFrame with date, steps, sleep, water
    """
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT date, steps, sleep_hours, water_intake
        FROM wellness_logs
        WHERE LOWER(patient_name) = LOWER(?)
        ORDER BY date
    """, (patient_name.strip(),))
    rows = c.fetchall()

    # goals (optional)
    c.execute("""
        SELECT steps_goal, sleep_goal, water_goal
        FROM wellness_goals
        WHERE LOWER(patient_name) = LOWER(?)
        LIMIT 1
    """, (patient_name.strip(),))
    g = c.fetchone()
    conn.close()

    if not rows:
        return None, None

    df = pd.DataFrame(rows, columns=["date", "steps", "sleep", "water"])
    df["date"] = pd.to_datetime(df["date"])
    df = df.sort_values("date")

    last7 = df.tail(7)
    stats = {
        "avg_steps": float(last7["steps"].mean()),
        "avg_sleep": float(last7["sleep"].mean()),
        "avg_water": float(last7["water"].mean()),
    }

    # streak of consecutive days with any wellness log (ending today)
    dates_set = set(df["date"].dt.date)
    today = date.today()
    streak = 0
    d = today
    while d in dates_set:
        streak += 1
        d = d - timedelta(days=1)
    stats["streak_days"] = streak

    # goal completion percentages (if goals exist)
    goal_completion = {}
    completion_values = []
    if g:
        steps_goal, sleep_goal, water_goal = g

        if steps_goal and steps_goal > 0:
            pct = min(stats["avg_steps"] / steps_goal, 1.0) * 100
            goal_completion["steps"] = pct
            completion_values.append(pct)

        if sleep_goal and sleep_goal > 0:
            pct = min(stats["avg_sleep"] / sleep_goal, 1.0) * 100
            goal_completion["sleep"] = pct
            completion_values.append(pct)

        if water_goal and water_goal > 0:
            pct = min(stats["avg_water"] / water_goal, 1.0) * 100
            goal_completion["water"] = pct
            completion_values.append(pct)

    stats["goal_completion"] = goal_completion
    stats["avg_goal_completion"] = (
        sum(completion_values) / len(completion_values)
        if completion_values
        else None
    )

    return stats, df


def create_dashboard_docx(
    patient_name,
    health_score,
    score_label,
    vitals,
    adherence_rate,
    wellness_stats,
    insights,
):
    """Create a DOCX dashboard summary for one patient."""
    doc = Document()
    doc.add_heading(f"Health Dashboard Summary - {patient_name}", level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    # Overall score
    doc.add_heading("Overall Health Score", level=2)
    doc.add_paragraph(f"Score: {health_score}/100 ({score_label})")
    doc.add_paragraph("")

    # Vitals
    doc.add_heading("Latest Vitals Snapshot", level=2)
    if vitals:
        doc.add_paragraph(f"Blood Pressure: {vitals['bp']}")
        doc.add_paragraph(f"Heart Rate: {vitals['heart_rate']} bpm")
        doc.add_paragraph(f"Temperature: {vitals['temperature']} °C")
        doc.add_paragraph(f"Risk Level: {vitals['risk']}")
        doc.add_paragraph(f"Symptoms: {vitals['symptoms']}")
    else:
        doc.add_paragraph("No vitals recorded yet.")
    doc.add_paragraph("")

    # Medication adherence
    doc.add_heading("Medication Adherence", level=2)
    if adherence_rate is not None:
        doc.add_paragraph(f"Overall adherence: {adherence_rate:.1f}%")
    else:
        doc.add_paragraph("No medication schedule found.")
    doc.add_paragraph("")

    # Wellness stats
    doc.add_heading("Wellness Trends (last 7 days)", level=2)
    if wellness_stats:
        doc.add_paragraph(
            f"Average steps: {wellness_stats['avg_steps']:.0f} steps/day"
        )
        doc.add_paragraph(
            f"Average sleep: {wellness_stats['avg_sleep']:.1f} hours/night"
        )
        doc.add_paragraph(
            f"Average water intake: {wellness_stats['avg_water']:.0f} ml/day"
        )
        doc.add_paragraph(
            f"Wellness streak: {wellness_stats['streak_days']} day(s) with logs in a row"
        )

        if wellness_stats["goal_completion"]:
            doc.add_paragraph("Goal completion:")
            for k, v in wellness_stats["goal_completion"].items():
                doc.add_paragraph(f"  • {k.capitalize()}: {v:.1f}%")
    else:
        doc.add_paragraph("No wellness logs found.")
    doc.add_paragraph("")

    # Insights
    doc.add_heading("AI Insights & Recommendations", level=2)
    if insights:
        for i, text in enumerate(insights, start=1):
            doc.add_paragraph(f"{i}. {text}")
    else:
        doc.add_paragraph("No insights generated.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =====================================================
# EXTRA HELPERS: WEEKLY PDF, MED CALENDAR, LIFESTYLE
# =====================================================

def create_weekly_pdf_report(
    patient_name, vitals, adherence_rate, wellness_stats, insights
):
    """Create a 1-page weekly progress PDF summary."""
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 40

    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, f"Weekly Progress Report - {patient_name}")
    y -= 30

    c.setFont("Helvetica", 10)
    c.drawString(40, y, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    y -= 25

    # Vitals
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Latest Vitals Snapshot")
    y -= 18
    c.setFont("Helvetica", 10)
    if vitals:
        lines = [
            f"Blood Pressure: {vitals['bp']}",
            f"Heart Rate: {vitals['heart_rate']} bpm",
            f"Temperature: {vitals['temperature']} °C",
            f"Risk Level: {vitals['risk']}",
            f"Symptoms: {vitals['symptoms']}",
        ]
    else:
        lines = ["No vitals recorded yet."]
    for line in lines:
        c.drawString(50, y, line)
        y -= 14

    y -= 10

    # Medication
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Medication Adherence")
    y -= 18
    c.setFont("Helvetica", 10)
    if adherence_rate is not None:
        c.drawString(50, y, f"Overall adherence this period: {adherence_rate:.1f}%")
    else:
        c.drawString(50, y, "No medication schedule found.")
    y -= 20

    # Wellness
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Wellness Summary (last 7 days)")
    y -= 18
    c.setFont("Helvetica", 10)
    if wellness_stats:
        w_lines = [
            f"Average steps: {wellness_stats['avg_steps']:.0f} steps/day",
            f"Average sleep: {wellness_stats['avg_sleep']:.1f} hours/night",
            f"Average water intake: {wellness_stats['avg_water']:.0f} ml/day",
            f"Wellness streak: {wellness_stats['streak_days']} day(s) in a row",
        ]
        for line in w_lines:
            c.drawString(50, y, line)
            y -= 14
    else:
        c.drawString(50, y, "No wellness logs recorded.")
        y -= 14

    y -= 10

    # Insights
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "AI Insights & Recommendations")
    y -= 18
    c.setFont("Helvetica", 10)
    if not insights:
        c.drawString(50, y, "No insights generated yet.")
    else:
        for text in insights:
            # wrap long insights roughly
            for chunk in [text[i:i+90] for i in range(0, len(text), 90)]:
                if y < 60:
                    c.showPage()
                    y = height - 60
                    c.setFont("Helvetica", 10)
                c.drawString(50, y, "- " + chunk)
                y -= 14

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


def build_medication_calendar_df(patient_name: str, month: int = None, year: int = None):
    """Return a calendar DataFrame for current month with emoji status."""
    patient_name = patient_name.strip()
    if month is None or year is None:
        today = date.today()
        month = today.month
        year = today.year

    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT id, medication_name, start_date, end_date
        FROM medications
        WHERE LOWER(patient_name) = LOWER(?)
    """, (patient_name,))
    meds = c.fetchall()

    if not meds:
        conn.close()
        return None

    med_ids = [m[0] for m in meds]
    placeholders = ",".join("?" for _ in med_ids)
    c.execute(f"""
        SELECT medication_id, log_date
        FROM medication_logs
        WHERE medication_id IN ({placeholders})
    """, med_ids)
    logs = c.fetchall()
    conn.close()

    logs_by_day = set()
    for _, log_date_str in logs:
        logs_by_day.add(datetime.strptime(log_date_str, "%Y-%m-%d").date())

    # helper: is any med active on day d?
    def has_schedule_on(d):
        for _, _, start_str, end_str in meds:
            s = datetime.strptime(start_str, "%Y-%m-%d").date()
            e = datetime.strptime(end_str, "%Y-%m-%d").date()
            if s <= d <= e:
                return True
        return False

    cal = calendar.Calendar(firstweekday=0)  # Monday
    month_days = cal.monthdatescalendar(year, month)

    today = date.today()
    table = []
    for week in month_days:
        row = []
        for d in week:
            if d.month != month:
                row.append("")
                continue
            if not has_schedule_on(d):
                row.append("—")  # no meds that day
            else:
                if d > today:
                    row.append("•")  # upcoming
                elif d in logs_by_day:
                    row.append("✅")  # taken (at least one dose)
                else:
                    row.append("❌")  # missed
        table.append(row)

    df = pd.DataFrame(
        table,
        columns=["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"],
    )
    return df


def generate_lifestyle_plan_for_patient(patient_name: str):
    """Rule-based lifestyle + diet + exercise suggestions."""
    patient_name = patient_name.strip()
    vitals = get_latest_vitals_for_dashboard(patient_name)
    summaries, adherence_rate, total_expected = overall_adherence_for_patient(
        patient_name
    )
    wellness_stats, _ = get_wellness_stats_for_dashboard(patient_name)

    plan = {
        "Daily Routine": [],
        "Diet": [],
        "Exercise": [],
        "Sleep & Stress": [],
        "Medication Habits": [],
    }

    # Risk-based
    if vitals and vitals["risk"] == "High":
        plan["Daily Routine"].append(
            "Avoid heavy physical or emotional stress; schedule a doctor visit soon."
        )
    elif vitals and vitals["risk"] == "Moderate":
        plan["Daily Routine"].append(
            "Monitor symptoms closely and avoid overexertion; keep a simple health diary."
        )
    else:
        plan["Daily Routine"].append(
            "Maintain regular routine with balanced work, movement and rest."
        )

    # Vitals-specific hints
    if vitals:
        if vitals["bp"]:
            plan["Daily Routine"].append(
                "Limit very salty and packaged foods to support healthy blood pressure."
            )
        try:
            temp_val = float(vitals["temperature"])
            if temp_val >= 37.5:
                plan["Daily Routine"].append(
                    "Stay indoors, rest more and increase fluids while fever lasts."
                )
        except Exception:
            pass

    # Wellness data
    if wellness_stats:
        if wellness_stats["avg_steps"] < 3000:
            plan["Exercise"].append(
                "Start with 10–15 minute walks 2–3 times a day, gradually target 5,000–8,000 steps."
            )
        else:
            plan["Exercise"].append(
                "Continue daily walking; add light stretching or yoga 3 times a week."
            )

        if wellness_stats["avg_sleep"] < 6:
            plan["Sleep & Stress"].append(
                "Aim for at least 7–8 hours of sleep; avoid screens 30 minutes before bed."
            )
        else:
            plan["Sleep & Stress"].append(
                "Keep a consistent sleep schedule with similar sleep and wake times daily."
            )

        if wellness_stats["avg_water"] < 1500:
            plan["Diet"].append(
                "Increase water intake; keep a bottle nearby and sip regularly through the day."
            )
        else:
            plan["Diet"].append(
                "Maintain 2–3 litres of water daily unless your doctor advised otherwise."
            )
    else:
        plan["Daily Routine"].append(
            "Start logging steps, sleep, and water so we can give more precise lifestyle guidance."
        )

    # Adherence
    if total_expected > 0 and adherence_rate is not None:
        if adherence_rate < 80:
            plan["Medication Habits"].append(
                "Set alarms or pill-box reminders so you do not miss doses; review medication schedule with your doctor."
            )
        else:
            plan["Medication Habits"].append(
                "Great job following medicines on time; continue using the same routine that works for you."
            )

    # General diet advice
    plan["Diet"].append(
        "Prefer home-cooked meals, whole grains, fruits and vegetables; limit deep-fried, very sugary and junk foods."
    )

    plan["Sleep & Stress"].append(
        "Practice 5–10 minutes of deep breathing, meditation or quiet time daily to reduce stress."
    )

    return plan


# =====================================================
# EXTRA ANALYTICS & AI HELPERS
# =====================================================

def get_risk_history(patient_name: str):
    """Return DataFrame of historical risk levels for a patient."""
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT id, risk_score
        FROM patients
        WHERE LOWER(name) = LOWER(?)
        ORDER BY id
    """, (patient_name.strip(),))
    rows = c.fetchall()
    conn.close()

    if not rows:
        return None

    # Map text risk to numeric for plotting
    risk_map = {"Low": 1, "Moderate": 2, "High": 3}
    visits = []
    scores = []
    labels = []
    for i, (pid, r) in enumerate(rows, start=1):
        visits.append(i)
        scores.append(risk_map.get(r, 0))
        labels.append(r or "Unknown")

    df = pd.DataFrame({
        "Visit": visits,
        "RiskScore": scores,
        "RiskLabel": labels,
    }).set_index("Visit")
    return df


def get_today_wellness(patient_name: str):
    """Return today's wellness log if present, else None."""
    today_str = date.today().strftime("%Y-%m-%d")
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT steps, sleep_hours, water_intake, mood
        FROM wellness_logs
        WHERE LOWER(patient_name) = LOWER(?)
          AND date = ?
        LIMIT 1
    """, (patient_name.strip(), today_str))
    row = c.fetchone()
    conn.close()
    if not row:
        return None
    return {
        "steps": row[0],
        "sleep": row[1],
        "water": row[2],
        "mood": row[3],
        "date": today_str,
    }


def build_daily_summary(patient_name: str):
    """Create a compact 'today summary' string + dict for dashboard card."""
    vitals = get_latest_vitals_for_dashboard(patient_name)
    summaries, adherence_rate, total_expected = overall_adherence_for_patient(
        patient_name
    )
    today_w = get_today_wellness(patient_name)

    summary_parts = []
    if vitals:
        summary_parts.append(f"Risk: {vitals['risk']}")
    if today_w:
        summary_parts.append(f"Steps {today_w['steps']}")
        summary_parts.append(f"Sleep {today_w['sleep']}h")
        summary_parts.append(f"Water {today_w['water']}ml")
        summary_parts.append(f"Mood: {today_w['mood']}")
    elif summaries:
        # fall back to adherence & averages
        summary_parts.append(f"Adherence {adherence_rate:.0f}%")
    else:
        summary_parts.append("No data recorded today")

    return " | ".join(summary_parts), {"vitals": vitals, "today": today_w,
                                       "adherence": adherence_rate,
                                       "total_expected": total_expected}


def generate_smart_recommendations(patient_name: str):
    """Short, action-oriented suggestions for 'today'."""
    tips = []

    vitals = get_latest_vitals_for_dashboard(patient_name)
    wellness_stats, _ = get_wellness_stats_for_dashboard(patient_name)
    _, adherence_rate, total_expected = overall_adherence_for_patient(patient_name)
    today_w = get_today_wellness(patient_name)

    # Based on risk
    if vitals and vitals.get("risk") == "High":
        tips.append(
            "Avoid heavy physical exertion today and watch for warning signs like chest pain or severe breathlessness."
        )
    elif vitals and vitals.get("risk") == "Moderate":
        tips.append(
            "Keep your day light; avoid skipping meals and stay hydrated while monitoring symptoms."
        )

    # Based on adherence
    if total_expected > 0 and adherence_rate is not None and adherence_rate < 80:
        tips.append(
            "Set a simple reminder (phone alarm) for your medicines so you don't miss doses today."
        )

    # Based on today's wellness
    if today_w:
        if today_w["steps"] < 3000:
            tips.append(
                "Try adding a 10–15 minute easy walk after lunch or dinner, if your doctor allows."
            )
        if today_w["sleep"] < 6:
            tips.append(
                "Plan for an earlier bedtime and avoid screens at least 30 minutes before sleep tonight."
            )
        if today_w["water"] < 1500:
            tips.append(
                "Keep a water bottle nearby and sip regularly so you reach your target by end of day."
            )
    elif wellness_stats:
        if wellness_stats["avg_steps"] < 3000:
            tips.append(
                "Start with small walks during the day; even 5 minutes each hour adds up."
            )

    if not tips:
        tips.append(
            "Continue your current routine and keep logging your vitals, wellness and medicines regularly."
        )

    return tips


def generate_indian_context_tips(patient_name: str):
    """Indian lifestyle-focused suggestions based on latest possible disease."""
    # Reuse latest patient data
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT symptoms, possible_diseases
        FROM patients
        WHERE LOWER(name) = LOWER(?)
        ORDER BY id DESC
        LIMIT 1
    """, (patient_name.strip(),))
    row = c.fetchone()
    conn.close()

    if not row:
        return ["No recent assessment found. Record at least one health check to see India-specific tips."]

    symptoms_text, diseases_text = row
    diseases = [d.strip() for d in diseases_text.split(",") if d.strip()] if diseases_text else []
    tips = []

    tips.append(
        "Prefer home-cooked Indian meals like dal, sabzi, idli, chapati, upma instead of deep-fried snacks such as samosa, pakoda and chips."
    )

    if any("Hypertension" in d or "BP" in d for d in diseases):
        tips.append(
            "For high BP, reduce very salty Indian foods like pickles, papad, packaged namkeen and bakery items."
        )

    if any("Diabetes" in d for d in diseases):
        tips.append(
            "For diabetes, limit sweets such as gulab jamun, rasgulla, barfi and sweet tea/coffee; choose plain roti, salads and controlled portions of fruit."
        )

    if any("Gastritis" in d or "Acidity" in d for d in diseases):
        tips.append(
            "For acidity, avoid very spicy oily curries, late-night heavy dinners and strong tea/coffee on an empty stomach."
        )

    if any("Asthma" in d or "COPD" in d for d in diseases):
        tips.append(
            "Asthma patients should avoid smoke from chulhas, incense, mosquito coils and bursting crackers; keep windows closed on high pollution days."
        )

    tips.append(
        "During monsoon and dengue season, avoid stagnant water around the house and use mosquito nets or repellents."
    )
    tips.append(
        "These are general Indian lifestyle tips. Always follow the specific advice of your treating doctor."
    )

    return tips


def answer_health_question(query: str, patient_name: str | None = None):
    """
    Very simple rule-based 'AI' for health Q&A.
    Gives general, safe information and always recommends seeing a doctor.
    """
    q = query.lower()
    answers = []

    # Link to medical info DB
    matches = lookup_med_info(query)
    if matches:
        for m in matches:
            answers.append(
                f"{m['name']}: {m['summary']} Common symptoms include: {m['common_symptoms']}"
            )

    # If asking about steps / activity / sleep / water
    if patient_name:
        if any(word in q for word in ["steps", "walk", "exercise", "activity"]):
            ws, _ = get_wellness_stats_for_dashboard(patient_name)
            if ws:
                answers.append(
                    f"Your recent average steps are about {ws['avg_steps']:.0f} per day. "
                    "If your doctor allows, many guidelines suggest 5,000–8,000 steps as a reasonable daily movement goal."
                )
        if any(word in q for word in ["sleep", "insomnia", "tired"]):
            ws, _ = get_wellness_stats_for_dashboard(patient_name)
            if ws:
                answers.append(
                    f"Your recent average sleep is around {ws['avg_sleep']:.1f} hours. "
                    "Most adults are advised to target 7–8 hours of good-quality sleep."
                )
        if any(word in q for word in ["water", "drink", "hydration"]):
            ws, _ = get_wellness_stats_for_dashboard(patient_name)
            if ws:
                answers.append(
                    f"Your recent average water intake is about {ws['avg_water']:.0f} ml. "
                    "A common target is around 2,000 ml/day unless your doctor has restricted fluids."
                )

    if not answers:
        answers.append(
            "I can give general health information, but I cannot confirm diagnoses or prescribe treatment. "
            "For serious or ongoing problems, please consult a qualified doctor."
        )

    answers.append(
        "This answer is for educational purposes only and does not replace professional medical advice."
    )
    return answers



# =====================================================
# MAIN HEADER
# =====================================================

st.markdown(
    "<h1 style='text-align:center; color:#333;'>🩺 Healthcare Monitoring AI Agent</h1>",
    unsafe_allow_html=True,
)
st.write(
    "<p style='text-align:center;'>Analyze symptoms, vitals & get health insights.</p>",
    unsafe_allow_html=True,
)

st.write("Enter your symptoms and vitals to check your health status.")


# =====================================================
# BASIC HEALTH ANALYZER
# =====================================================

name = st.text_input("Full Name")
age = st.number_input("Age", min_value=1, max_value=120)
symptoms = st.text_area("Describe your symptoms")
bp = st.text_input("Blood Pressure (e.g., 120/80)")
heart_rate = st.text_input("Heart Rate (bpm)")
temperature = st.text_input("Body Temperature (°C)")

if st.button("Analyze Health"):
    if name and symptoms and bp and heart_rate and temperature:
        ok, val_errors = validate_vitals_inputs(bp, heart_rate, temperature)
        if not ok:
            for msg in val_errors:
                st.error(msg)
        else:
            try:
                extracted = extract_symptoms(symptoms)
                risk, issues = analyze_vitals(bp, heart_rate, temperature)
                diseases = match_disease(extracted)  # could be list[str] or list[dict]

                st.subheader("🧪 Health Analysis Result")

                if risk == "Low":
                    st.success("🟢 Risk Level: LOW")
                elif risk == "Moderate":
                    st.warning("🟡 Risk Level: MODERATE")
                else:
                    st.error("🔴 Risk Level: HIGH")

                st.write("### Issues detected:")
                for i in issues:
                    st.write(f"- {i}")

                st.write("### Possible Diseases:")
                if diseases:
                    for d in diseases:
                        st.write(f"- {disease_name(d)}")
                else:
                    st.write("- Not enough information to guess a condition.")

                possible = ", ".join(disease_name(d) for d in diseases) if diseases else ""

                conn = sqlite3.connect("healthcare.db")
                c = conn.cursor()
                c.execute(
                    """
                    INSERT INTO patients (
                        name, age, symptoms, bp, heart_rate, temperature, risk_score, possible_diseases
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (name, age, symptoms, bp, heart_rate, temperature, risk, possible),
                )
                conn.commit()
                conn.close()

                st.success("✔ Health data stored and analyzed successfully!")
            except Exception as e:
                st.error(f"Something went wrong while analyzing or saving data: {e}")
    else:
        st.error("Please fill all fields.")



# =====================================================
# ADVANCED AI SYMPTOM HELPER
# =====================================================

st.markdown("---")
st.header("🤖 AI Symptom Helper (Advanced Version)")

user_question = st.text_input(
    "Describe your symptoms (e.g., I have fever, cough, body pain):"
)

if st.button("Analyze Symptoms"):
    if not user_question.strip():
        st.error("Please type your symptoms first.")
    else:
        extracted = extract_symptoms(user_question)

        if not extracted:
            st.warning(
                "No known symptoms detected. Try using simple words like 'fever', "
                "'cough', 'headache'."
            )
        else:
            st.success(f"Detected symptoms: {', '.join(extracted)}")

            # Symptom explanations (rule-based)
            explanations = explain_symptoms(extracted)
            if explanations:
                st.subheader("🔍 Symptom Explanations")
                for e in explanations:
                    st.write(f"- {e}")

            # Disease ranking (KB-based)
            ranked = rank_diseases(extracted)
            st.subheader("🩺 Possible Conditions (Ranked)")

            if not ranked:
                st.info("No strong matches found in the disease knowledge base.")
            else:
                top_n = ranked[:3]
                for idx, item in enumerate(top_n, start=1):
                    st.markdown(
                        f"""
**{idx}. {item['disease']}**

- Matched symptoms: {', '.join(item['matched_symptoms'])}  
- Match score: **{item['score']}** / {len(DISEASE_KB[item['disease']])}
                        """
                    )

            # Doctor type recommendation (from doctor_engine)
            st.subheader("👨‍⚕️ Recommended Doctor Type")
            top_disease_name, specialities = suggest_specialities(ranked)
            if top_disease_name:
                st.write(f"Top suspected condition: **{top_disease_name}**")
            if specialities:
                st.write("Suggested specialists:")
                for s in specialities:
                    st.write(f"- {s}")



            # Follow-up plan
            st.subheader("⏰ Follow-up Recommendation")
            symptom_risk = estimate_risk_from_symptoms(extracted)
            f_date, f_msg = get_followup_plan(symptom_risk)
            st.write(f"Risk Level (symptom-based): **{symptom_risk}**")
            st.write(f"Suggested follow-up date: **{f_date}**")
            st.write(f"{f_msg}")

            # Care tips
            st.subheader("💡 Care Suggestions")
            care_tips = get_care_tips(extracted, top_disease_name)
            for tip in care_tips:
                st.write(f"• {tip}")


# =====================================================
# SIMPLE MEDICAL INFORMATION KB (FOR LOOKUP)
# =====================================================

MED_INFO_DB = [
    {
        "name": "Hypertension (High Blood Pressure)",
        "keywords": ["hypertension", "high bp", "high blood pressure", "bp"],
        "summary": "A condition where the force of blood against artery walls is consistently too high.",
        "common_symptoms": "Often no symptoms; sometimes headache, shortness of breath, nosebleeds in severe cases.",
        "self_care": "Reduce salt intake, maintain healthy weight, be physically active, avoid smoking and limit alcohol as advised by a doctor.",
        "sources": [
            "WHO – Hypertension fact sheet",
            "CDC – High Blood Pressure information page"
        ],
    },
    {
        "name": "Type 2 Diabetes",
        "keywords": ["diabetes", "sugar", "high sugar"],
        "summary": "A long-term condition where the body does not use insulin properly, leading to high blood sugar.",
        "common_symptoms": "Increased thirst, frequent urination, tiredness, slow-healing wounds, blurred vision.",
        "self_care": "Follow prescribed medicines, monitor blood sugar, choose balanced meals, and stay active as advised by your doctor.",
        "sources": [
            "WHO – Diabetes fact sheet",
            "International Diabetes Federation – Type 2 Diabetes overview"
        ],
    },
    {
        "name": "Asthma",
        "keywords": ["asthma", "wheezing", "breathlessness"],
        "summary": "A chronic condition where airways become inflamed and narrow, making breathing difficult.",
        "common_symptoms": "Wheezing, cough, chest tightness, shortness of breath, especially at night or early morning.",
        "self_care": "Use inhalers as prescribed, avoid known triggers (dust, smoke, pollution), and follow an asthma action plan.",
        "sources": [
            "GINA – Global Initiative for Asthma",
            "WHO – Asthma fact sheet"
        ],
    },
    {
        "name": "Migraine",
        "keywords": ["migraine", "severe headache", "one-sided headache"],
        "summary": "A type of headache that can cause intense, throbbing pain, often on one side of the head.",
        "common_symptoms": "Severe headache, nausea, sensitivity to light or sound, sometimes visual disturbances (aura).",
        "self_care": "Rest in a dark quiet room, stay hydrated, avoid known triggers like certain foods, lack of sleep or stress.",
        "sources": [
            "Mayo Clinic – Migraine overview",
            "NHS – Migraine information"
        ],
    },
    {
        "name": "Depression",
        "keywords": ["depression", "low mood", "sadness"],
        "summary": "A common mental health condition with persistent sadness and loss of interest or pleasure.",
        "common_symptoms": "Low mood, loss of interest, changes in sleep or appetite, feelings of worthlessness, difficulty concentrating.",
        "self_care": "Talk to a trusted person, maintain regular routine, gentle physical activity if possible, and seek professional help early.",
        "sources": [
            "WHO – Depression fact sheet",
            "NIMH – Depression information page"
        ],
    },
]


def lookup_med_info(query: str):
    """Return list of KB entries matching the query."""
    q = query.lower().strip()
    results = []
    for item in MED_INFO_DB:
        if item["name"].lower() in q:
            results.append(item)
            continue
        if any(kw in q for kw in item["keywords"]):
            results.append(item)
    return results


# =====================================================
# MEDICAL INFORMATION LOOKUP (WITH SOURCES)
# =====================================================

st.markdown("---")
st.header("📚 Medical Information Lookup")

st.caption(
    "This section provides general health information from well-known medical organisations. "
    "It is for education only and does not replace a doctor’s advice."
)

info_query = st.text_input(
    "Search a condition (e.g., diabetes, asthma, high BP, migraine):"
)

if st.button("Search Medical Information"):
    if not info_query.strip():
        st.error("Please type a disease or condition name.")
    else:
        matches = lookup_med_info(info_query)
        if not matches:
            st.info(
                "This condition is not in the built-in knowledge base yet. "
                "Please consult a trusted medical website or your doctor."
            )
        else:
            for item in matches:
                st.subheader(item["name"])
                st.write(f"**Overview:** {item['summary']}")
                st.write(f"**Common symptoms:** {item['common_symptoms']}")
                st.write(f"**Self-care tips (general):** {item['self_care']}")
                st.write("**Sources (information websites):**")
                for src in item["sources"]:
                    st.write(f"- {src}")



# =====================================================
# VIEW & EXPORT STORED HEALTH RECORDS
# =====================================================

st.markdown("---")
st.header("📋 View & Export Stored Health Records")

if st.button("Load Patient Records"):
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute(
        """
        SELECT name, age, symptoms, bp, heart_rate, temperature, risk_score, possible_diseases
        FROM patients
        """
    )
    rows = c.fetchall()
    conn.close()

    if not rows:
        st.info("No records found.")
    else:
        st.subheader("Saved Records")
        patients = [row_to_patient(r) for r in rows]

        for p in patients:
            st.write(f"**Name:** {p['name']}")
            st.write(f"**Age:** {p['age']}")
            st.write(f"**Symptoms:** {p['symptoms']}")
            st.write(f"**BP:** {p['bp']}")
            st.write(f"**Heart Rate:** {p['heart_rate']}")
            st.write(f"**Temperature:** {p['temperature']}")
            st.write(f"**Risk:** {p['risk_score']}")
            st.write(f"**Possible Diseases:** {p['possible_diseases']}")
            st.markdown("---")

        st.subheader("Export Reports")

        selected_index = st.selectbox(
            "Select a patient to export:",
            options=list(range(len(patients))),
            format_func=lambda i: f"{patients[i]['name']} (Age {patients[i]['age']})",
        )
        selected_patient = patients[selected_index]

        col1, col2 = st.columns(2)

        with col1:
            docx_buf = create_docx_for_patient(selected_patient)
            st.download_button(
                label="⬇ Download selected as DOCX",
                data=docx_buf,
                file_name=f"{selected_patient['name'].replace(' ', '_')}_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        with col2:
            pdf_buf = create_pdf_for_patient(selected_patient)
            st.download_button(
                label="⬇ Download selected as PDF",
                data=pdf_buf,
                file_name=f"{selected_patient['name'].replace(' ', '_')}_report.pdf",
                mime="application/pdf",
            )

        st.markdown("### Export All Records")

        all_docx_buf = create_docx_for_all(patients)
        st.download_button(
            label="⬇ Download ALL as DOCX",
            data=all_docx_buf,
            file_name="all_patients_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        all_pdf_buf = create_pdf_for_all(patients)
        st.download_button(
            label="⬇ Download ALL as PDF",
            data=all_pdf_buf,
            file_name="all_patients_report.pdf",
            mime="application/pdf",
        )

def get_all_patient_names():
    """Return list of distinct patient names from the DB."""
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("SELECT DISTINCT name FROM patients ORDER BY name")
    rows = c.fetchall()
    conn.close()
    return [r[0] for r in rows]


# =====================================================
# PATIENT COMMUNICATION TEMPLATES
# =====================================================

st.markdown("---")
st.header("📨 Patient Communication Templates")


def get_latest_patient_for_msg():
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute(
        """
        SELECT name, risk_score, symptoms, possible_diseases
        FROM patients
        ORDER BY id DESC
        LIMIT 1
        """
    )
    row = c.fetchone()
    conn.close()

    if row:
        return {
            "name": row[0],
            "risk": row[1],
            "symptoms": row[2],
            "diseases": row[3],
        }
    return None


def msg_followup(p):
    return f"""
Dear {p['name']},

Based on your recent health evaluation, your overall risk level was **{p['risk']}**.

Recorded symptoms:
{p['symptoms']}

Possible conditions:
{p['diseases']}

👉 Please monitor your symptoms for any changes.  
👉 Stay hydrated, take enough rest, and follow your medication schedule.  
👉 If your symptoms worsen or new symptoms appear, seek medical care immediately.

Wishing you good health,  
Healthcare Monitoring AI Agent
"""


def msg_emergency(p):
    return f"""
Dear {p['name']},

Our recent assessment shows your risk level as **{p['risk']}**.

Symptoms noted:
{p['symptoms']}

Possible serious conditions:
{p['diseases']}

⚠ We recommend that you **do not ignore these symptoms**.

• If you experience chest pain, severe breathlessness, confusion, or fainting,  
  please visit the nearest emergency department immediately.  

Stay safe,  
Healthcare Monitoring AI Agent
"""


def msg_recovery_reminder(p):
    return f"""
Hello {p['name']},

This is a gentle reminder to support your recovery.

Your last check showed risk level: **{p['risk']}**  
Symptoms: {p['symptoms']}

Please remember:
• Take prescribed medicines on time.  
• Do not skip doses without asking your doctor.  
• Get enough rest and sleep.  

Get well soon!  
Healthcare Monitoring AI Agent
"""


def msg_lifestyle_diet(p):
    return f"""
Hi {p['name']},

Here are some lifestyle and diet suggestions based on your recent checkup.

Symptoms: {p['symptoms']}  
Possible conditions: {p['diseases']}

General tips:
• Drink 2–3 litres of water a day (unless your doctor advised restriction).  
• Include fruits, vegetables, and whole grains in your meals.  
• Avoid very oily, spicy, or junk food.  
• Limit sugary drinks and smoking/alcohol (if applicable).  

These are general suggestions and do not replace a doctor's advice.

With care,  
Healthcare Monitoring AI Agent
"""


def msg_appointment_confirmation(p, when_text):
    return f"""
Dear {p['name']},

Your medical appointment has been booked.

🗓 Appointment details:
• Patient: {p['name']}  
• Date & Time: {when_text}

Recent assessment:
• Risk level: {p['risk']}  
• Symptoms: {p['symptoms']}  
• Possible conditions: {p['diseases']}

Please reach the clinic/hospital 10–15 minutes early and
carry your previous medical reports and medicine list.

Regards,  
Healthcare Monitoring AI Agent
"""


template_type = st.selectbox(
    "Choose a template to generate:",
    [
        "Follow-up Advice",
        "Emergency Alert",
        "Recovery / Medicine Reminder",
        "Lifestyle & Diet Suggestions",
        "Appointment Confirmation",
    ],
)

appt_datetime_text = None
if template_type == "Appointment Confirmation":
    c1, c2 = st.columns(2)
    with c1:
        appt_date_msg = st.date_input("Select appointment date")
    with c2:
        appt_time_msg = st.time_input("Select appointment time")
    appt_datetime_text = (
        f"{appt_date_msg.strftime('%Y-%m-%d')} at {appt_time_msg.strftime('%H:%M')}"
    )

if st.button("Generate Message"):
    p = get_latest_patient_for_msg()
    if not p:
        st.warning("⚠ No patient data found. Please analyze at least one patient first.")
    else:
        if template_type == "Follow-up Advice":
            message = msg_followup(p)
        elif template_type == "Emergency Alert":
            message = msg_emergency(p)
        elif template_type == "Recovery / Medicine Reminder":
            message = msg_recovery_reminder(p)
        elif template_type == "Lifestyle & Diet Suggestions":
            message = msg_lifestyle_diet(p)
        else:
            message = msg_appointment_confirmation(p, appt_datetime_text)

        st.subheader("Generated Message")
        st.text_area("You can copy this message:", message, height=260)


# =====================================================
# DOCTOR SPECIALITY MAPPING FOR SCHEDULER
# =====================================================

def scheduler_speciality(disease):
    mapping = {
        "Hypertension": "Cardiologist",
        "Hypertension (High BP)": "Cardiologist",
        "Diabetes": "Endocrinologist",
        "Asthma": "Pulmonologist",
        "Migraine": "Neurologist",
        "Flu": "General Physician",
        "COVID-19": "Pulmonologist",
        "Gastritis": "Gastroenterologist",
        "UTI": "Urologist",
        "Arthritis": "Rheumatologist",
        "Skin Allergy": "Dermatologist",
        "Depression": "Psychiatrist",
        "Anxiety": "Psychiatrist",
        "Thyroid Disorder": "Endocrinologist",
        "Dengue": "General Physician",
        "Viral Fever": "General Physician",
        "Heart Disease": "Cardiologist",
        "Kidney Stone": "Urologist",
        "PCOS": "Gynecologist",
        "Pregnancy": "Gynecologist",
        "Ear Infection": "ENT Specialist",
        "Common Cold": "General Physician",
        "Acidity": "Gastroenterologist",
    }
    return mapping.get(disease, "General Physician")


# 19-speciality doctor DB (for scheduler)
DOCTOR_DB = {
    "General Physician": [
        {"name": "Dr. Anil Kumar", "experience": 12},
        {"name": "Dr. Swetha Rao", "experience": 8},
        {"name": "Dr. Rajesh Singh", "experience": 15},
    ],
    "Cardiologist": [
        {"name": "Dr. Meera Nair", "experience": 14},
        {"name": "Dr. Ashok Patil", "experience": 10},
        {"name": "Dr. Virat Shah", "experience": 18},
    ],
    "Neurologist": [
        {"name": "Dr. Sahana Iyer", "experience": 9},
        {"name": "Dr. Rishi Menon", "experience": 13},
        {"name": "Dr. Neha Bhat", "experience": 11},
    ],
    "Dermatologist": [
        {"name": "Dr. Kavya Shetty", "experience": 7},
        {"name": "Dr. Rohan Kapoor", "experience": 12},
        {"name": "Dr. Trisha Jain", "experience": 10},
    ],
    "ENT Specialist": [
        {"name": "Dr. Jayanth Kulkarni", "experience": 10},
        {"name": "Dr. Pooja Sharma", "experience": 6},
        {"name": "Dr. Raghav Mehta", "experience": 9},
    ],
    "Orthopedic": [
        {"name": "Dr. Manoj Reddy", "experience": 16},
        {"name": "Dr. Divya Joshi", "experience": 11},
        {"name": "Dr. Jay Patel", "experience": 14},
    ],
    "Gastroenterologist": [
        {"name": "Dr. Nikita Das", "experience": 10},
        {"name": "Dr. Karthik S", "experience": 12},
        {"name": "Dr. Mahesh Agarwal", "experience": 15},
    ],
    "Psychiatrist": [
        {"name": "Dr. Ananya Rao", "experience": 8},
        {"name": "Dr. Varun Khanna", "experience": 10},
        {"name": "Dr. Radhika Pillai", "experience": 11},
    ],
    "Pulmonologist": [
        {"name": "Dr. Harish Prasad", "experience": 13},
        {"name": "Dr. Sneha M", "experience": 9},
        {"name": "Dr. Ajay Verma", "experience": 7},
    ],
    "Endocrinologist": [
        {"name": "Dr. Sonia Mishra", "experience": 12},
        {"name": "Dr. Farhan Khan", "experience": 15},
        {"name": "Dr. Ritu Malhotra", "experience": 9},
    ],
    "Diabetologist": [
        {"name": "Dr. Sunil Shetty", "experience": 10},
        {"name": "Dr. Priya Desai", "experience": 11},
        {"name": "Dr. Mohan Rao", "experience": 7},
    ],
    "Nephrologist": [
        {"name": "Dr. Kavita R", "experience": 14},
        {"name": "Dr. Naveen P", "experience": 10},
        {"name": "Dr. Aarav Kumar", "experience": 8},
    ],
    "Urologist": [
        {"name": "Dr. Prasad Gowda", "experience": 13},
        {"name": "Dr. Divya Patel", "experience": 7},
        {"name": "Dr. Jayant S", "experience": 11},
    ],
    "Gynecologist": [
        {"name": "Dr. Aishwarya N", "experience": 12},
        {"name": "Dr. Farida Khan", "experience": 15},
        {"name": "Dr. Manasa M", "experience": 9},
    ],
    "Pediatrician": [
        {"name": "Dr. Nandini R", "experience": 10},
        {"name": "Dr. Aditya S", "experience": 8},
        {"name": "Dr. Mehul Jain", "experience": 13},
    ],
    "Oncologist": [
        {"name": "Dr. Ramesh N", "experience": 16},
        {"name": "Dr. Shruti Arora", "experience": 12},
        {"name": "Dr. Vivek Rao", "experience": 14},
    ],
    "Rheumatologist": [
        {"name": "Dr. Sanjana Rao", "experience": 11},
        {"name": "Dr. Ajith K", "experience": 9},
        {"name": "Dr. Preeti Narang", "experience": 12},
    ],
    "Hematologist": [
        {"name": "Dr. Anoop S", "experience": 13},
        {"name": "Dr. Megha R", "experience": 7},
        {"name": "Dr. Suraj R", "experience": 10},
    ],
    "Ophthalmologist": [
        {"name": "Dr. Kiran Shetty", "experience": 15},
        {"name": "Dr. Nisha Kapoor", "experience": 11},
        {"name": "Dr. Rohit Bhat", "experience": 9},
    ],
}

# random availability table — ALL days, same 3 slots
DAYS = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat"]
SLOTS = ["10:00 AM", "1:00 PM", "4:00 PM"]

DOCTOR_AVAILABILITY = {}
for speciality, doctors in DOCTOR_DB.items():
    DOCTOR_AVAILABILITY[speciality] = {}
    for d in doctors:
        name_doc = d["name"]
        DOCTOR_AVAILABILITY[speciality][name_doc] = {
            day: SLOTS for day in DAYS
        }


def weekday_from_date(dt):
    return dt.strftime("%a")  # Mon/Tue/...

# =====================================================
# SIMPLE MEDICAL INFORMATION KB (FOR LOOKUP)
# =====================================================

MED_INFO_DB = [
    {
        "name": "Hypertension (High Blood Pressure)",
        "keywords": ["hypertension", "high bp", "high blood pressure", "bp"],
        "summary": "A condition where the force of blood against artery walls is consistently too high.",
        "common_symptoms": "Often no symptoms; sometimes headache, shortness of breath, nosebleeds in severe cases.",
        "self_care": "Reduce salt intake, maintain healthy weight, be physically active, avoid smoking and limit alcohol as advised by a doctor.",
        "sources": [
            "WHO – Hypertension fact sheet",
            "CDC – High Blood Pressure information page"
        ],
    },
    {
        "name": "Type 2 Diabetes",
        "keywords": ["diabetes", "sugar", "high sugar"],
        "summary": "A long-term condition where the body does not use insulin properly, leading to high blood sugar.",
        "common_symptoms": "Increased thirst, frequent urination, tiredness, slow-healing wounds, blurred vision.",
        "self_care": "Follow prescribed medicines, monitor blood sugar, choose balanced meals, and stay active as advised by your doctor.",
        "sources": [
            "WHO – Diabetes fact sheet",
            "International Diabetes Federation – Type 2 Diabetes overview"
        ],
    },
    {
        "name": "Asthma",
        "keywords": ["asthma", "wheezing", "breathlessness"],
        "summary": "A chronic condition where airways become inflamed and narrow, making breathing difficult.",
        "common_symptoms": "Wheezing, cough, chest tightness, shortness of breath, especially at night or early morning.",
        "self_care": "Use inhalers as prescribed, avoid known triggers (dust, smoke, pollution), and follow an asthma action plan.",
        "sources": [
            "GINA – Global Initiative for Asthma",
            "WHO – Asthma fact sheet"
        ],
    },
    {
        "name": "Migraine",
        "keywords": ["migraine", "severe headache", "one-sided headache"],
        "summary": "A type of headache that can cause intense, throbbing pain, often on one side of the head.",
        "common_symptoms": "Severe headache, nausea, sensitivity to light or sound, sometimes visual disturbances (aura).",
        "self_care": "Rest in a dark quiet room, stay hydrated, avoid known triggers like certain foods, lack of sleep or stress.",
        "sources": [
            "Mayo Clinic – Migraine overview",
            "NHS – Migraine information"
        ],
    },
    {
        "name": "Depression",
        "keywords": ["depression", "low mood", "sadness"],
        "summary": "A common mental health condition with persistent sadness and loss of interest or pleasure.",
        "common_symptoms": "Low mood, loss of interest, changes in sleep or appetite, feelings of worthlessness, difficulty concentrating.",
        "self_care": "Talk to a trusted person, maintain regular routine, gentle physical activity if possible, and seek professional help early.",
        "sources": [
            "WHO – Depression fact sheet",
            "NIMH – Depression information page"
        ],
    },
]

def lookup_med_info(query: str):
    q = query.lower().strip()
    results = []
    for item in MED_INFO_DB:
        if item["name"].lower() in q:
            results.append(item)
            continue
        if any(kw in q for kw in item["keywords"]):
            results.append(item)
    return results


# =====================================================
# SMART APPOINTMENT SCHEDULER (SINGLE, CLEAN VERSION)
# =====================================================

st.markdown("---")
st.header("📅 Smart Appointment Scheduler (AI Recommended)")

patient_names = get_all_patient_names()
if patient_names:
    appt_patient_name = st.selectbox(
        "Select patient for appointment (from existing records):",
        patient_names,
    )
else:
    st.info("No patient records found yet. Please analyze at least one patient first.")
    appt_patient_name = ""


# State for scheduler
if "sched_speciality" not in st.session_state:
    st.session_state.sched_speciality = None
if "sched_disease" not in st.session_state:
    st.session_state.sched_disease = None
if "sched_doctor" not in st.session_state:
    st.session_state.sched_doctor = None
if "sched_date" not in st.session_state:
    st.session_state.sched_date = None

# STEP 1 – infer disease & speciality from last analysis
if st.button("Find Recommended Doctor"):
    if not appt_patient_name.strip():
        st.error("Please enter patient name.")
    else:
        latest = fetch_latest_patient(appt_patient_name)
        if latest is None:
            st.warning("⚠ No previous health analysis found for this patient.")
        else:
            stored_symptoms, stored_risk = latest
            extracted = extract_symptoms(stored_symptoms)
            ranked = rank_diseases(extracted)

            if not ranked:
                st.warning("Cannot detect disease. Try adding more symptoms.")
            else:
                top_disease = ranked[0]["disease"]
                speciality = scheduler_speciality(top_disease)

                st.session_state.sched_disease = top_disease
                st.session_state.sched_speciality = speciality

                st.success(f"Detected Condition: **{top_disease}**")
                st.info(f"Recommended Specialist: **{speciality}**")

# STEP 2 – doctor selection
if st.session_state.sched_speciality:
    speciality = st.session_state.sched_speciality
    st.subheader(f"👨‍⚕️ Available {speciality}s")

    doctors = DOCTOR_DB.get(speciality, [])
    doctor_names = [d["name"] for d in doctors]

    if doctor_names:
        st.session_state.sched_doctor = st.selectbox(
            "Select Doctor:",
            doctor_names,
        )
    else:
        st.warning("No doctors configured for this speciality.")

# STEP 3 – date + slots
if st.session_state.sched_doctor:
    st.subheader("📆 Select Appointment Date")
    appt_date = st.date_input("Choose a date")
    st.session_state.sched_date = appt_date

    if appt_date:
        weekday = weekday_from_date(appt_date)
        doctor = st.session_state.sched_doctor

        st.write(f"Selected Day: **{weekday}**")

        slots_map = DOCTOR_AVAILABILITY[
            st.session_state.sched_speciality
        ].get(doctor, {})

        if weekday in slots_map:
            slots = slots_map[weekday]
            st.subheader("🕒 Available Slots")
            selected_slot = st.radio("Choose a time slot:", slots)

            # STEP 4 – confirm + store in DB
            if st.button("Confirm Appointment"):
                conn = sqlite3.connect("healthcare.db")
                c = conn.cursor()
                c.execute(
                    """
                    INSERT INTO appointments (doctor, speciality, patient, date, time)
                    VALUES (?, ?, ?, ?, ?)
                    """,
                    (
                        doctor,
                        st.session_state.sched_speciality,
                        appt_patient_name,
                        appt_date.strftime("%Y-%m-%d"),
                        selected_slot,
                    ),
                )
                conn.commit()
                conn.close()

                st.success(
                    f"✔ Appointment Confirmed!\n\n"
                    f"📌 Patient: **{appt_patient_name}**\n"
                    f"👨‍⚕️ Doctor: **{doctor}**\n"
                    f"🏥 Speciality: **{st.session_state.sched_speciality}**\n"
                    f"📅 Date: **{appt_date.strftime('%Y-%m-%d')}**\n"
                    f"⏰ Time: **{selected_slot}**"
                )
        else:
            st.warning("❌ This doctor is not available on the selected day.")


# =====================================================
# MEDICATION ADHERENCE TRACKER
# =====================================================

def compute_adherence(med_row, logs):
    """
    med_row: row from medications table
    logs: list of rows from medication_logs for this medication
    """
    # unpack
    _, _, _, _, freq, times_str, start, end = med_row
    times_list = [t.strip() for t in times_str.split(",") if t.strip()]
    doses_per_day = len(times_list)

    start_d = datetime.strptime(start, "%Y-%m-%d").date()
    end_d = datetime.strptime(end, "%Y-%m-%d").date()
    today = date.today()

    # Only count until today (or end date, whichever is smaller)
    last_day = min(today, end_d)
    if last_day < start_d:
        return 0, 0, 0.0  # Not started yet

    days = (last_day - start_d).days + 1
    expected = days * doses_per_day
    taken = len(logs)
    adherence = (taken / expected * 100) if expected > 0 else 0.0
    return expected, taken, adherence


def overall_adherence_for_patient(patient_name: str):
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT id, patient_name, medication_name, dosage,
               frequency, time_of_day, start_date, end_date
        FROM medications
        WHERE LOWER(patient_name) = LOWER(?)
    """, (patient_name.strip(),))
    meds = c.fetchall()

    if not meds:
        conn.close()
        return [], 0.0, 0

    summaries = []
    total_expected = 0
    total_taken = 0

    for med in meds:
        med_id = med[0]
        c.execute("""
            SELECT id, medication_id, log_date, time_of_day, status
            FROM medication_logs
            WHERE medication_id = ?
        """, (med_id,))
        logs = c.fetchall()
        expected, taken, adherence = compute_adherence(med, logs)
        total_expected += expected
        total_taken += taken
        summaries.append(
            {
                "name": med[2],
                "dosage": med[3],
                "period": f"{med[6]} → {med[7]}",
                "expected": expected,
                "taken": taken,
                "adherence": adherence,
            }
        )

    conn.close()
    overall = (total_taken / total_expected * 100) if total_expected > 0 else 0.0
    return summaries, overall, total_expected


def upsert_wellness_goal(patient_name, steps_goal, sleep_goal, water_goal):
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT id FROM wellness_goals
        WHERE LOWER(patient_name) = LOWER(?)
    """, (patient_name.strip(),))
    row = c.fetchone()

    if row:
        c.execute("""
            UPDATE wellness_goals
            SET steps_goal = ?, sleep_goal = ?, water_goal = ?
            WHERE id = ?
        """, (steps_goal, sleep_goal, water_goal, row[0]))
    else:
        c.execute("""
            INSERT INTO wellness_goals (patient_name, steps_goal, sleep_goal, water_goal)
            VALUES (?, ?, ?, ?)
        """, (patient_name, steps_goal, sleep_goal, water_goal))

    conn.commit()
    conn.close()


def log_wellness_entry(patient_name, steps, sleep_hours, water_intake, mood, log_date=None):
    if log_date is None:
        log_date = date.today().strftime("%Y-%m-%d")
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        INSERT INTO wellness_logs (patient_name, steps, sleep_hours, water_intake, mood, date)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (patient_name, steps, sleep_hours, water_intake, mood, log_date))
    conn.commit()
    conn.close()


def generate_insights(patient_name: str):
    patient_name = patient_name.strip()
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()

    # latest risk from vitals
    c.execute("""
        SELECT risk_score
        FROM patients
        WHERE LOWER(name) = LOWER(?)
        ORDER BY id DESC
        LIMIT 1
    """, (patient_name,))
    pat_row = c.fetchone()

    # medication adherence
    summaries, overall_rate, total_expected = overall_adherence_for_patient(patient_name)

    # recent wellness (last 7 entries)
    c.execute("""
        SELECT steps, sleep_hours, water_intake
        FROM wellness_logs
        WHERE LOWER(patient_name) = LOWER(?)
        ORDER BY date DESC
        LIMIT 7
    """, (patient_name,))
    w_rows = c.fetchall()

    conn.close()

    insights = []

    if pat_row:
        risk = pat_row[0]
        if risk == "High":
            insights.append(
                "Risk level is HIGH based on vitals. Regular doctor follow-up is important."
            )
        elif risk == "Moderate":
            insights.append(
                "Risk level is MODERATE. Monitor symptoms and keep healthy habits."
            )
        elif risk == "Low":
            insights.append("Risk level is LOW. Maintain your current lifestyle.")

    if total_expected > 0:
        if overall_rate >= 90:
            insights.append(
                "Medication adherence is excellent (≥ 90%). Keep following the prescription on time."
            )
        elif overall_rate >= 80:
            insights.append(
                "Medication adherence is good, but try not to miss doses to stay on track."
            )
        else:
            insights.append(
                "Medication adherence is low (< 80%). Missing medicines may reduce treatment effectiveness; talk to your doctor."
            )

    if w_rows:
        avg_steps = sum(r[0] for r in w_rows) / len(w_rows)
        avg_sleep = sum(r[1] for r in w_rows) / len(w_rows)
        avg_water = sum(r[2] for r in w_rows) / len(w_rows)

        insights.append(
            f"Average daily steps (last week): {avg_steps:.0f}. Try to reach at least 5,000–8,000 steps if your doctor allows."
        )
        insights.append(
            f"Average sleep: {avg_sleep:.1f} hours. Aim for 7–8 hours of sleep per night."
        )
        insights.append(
            f"Average water intake: {avg_water:.0f} ml. Around 2,000 ml/day is a common target unless your doctor advises otherwise."
        )

    if not insights:
        insights.append(
            "Not enough data yet. Add vitals, medications, and wellness logs to see insights."
        )

    return insights


st.markdown("---")
st.header("💊 Medication Adherence Tracker")

med_patient = st.text_input("Patient Name (for medication tracking)")
med_name = st.text_input("Medication Name")
med_dosage = st.text_input("Dosage (e.g., 500 mg)")
med_freq = st.selectbox("Frequency", ["Once a day", "Twice a day", "Thrice a day"])
med_times = st.multiselect("Time(s) of Day", ["Morning", "Afternoon", "Evening", "Night"])
med_start = st.date_input("Start Date")
med_end = st.date_input("End Date")

if st.button("Add Medication"):
    ok_basic = (med_patient and med_name and med_dosage and med_times)
    if not ok_basic:
        st.error("Please fill patient, medicine, dosage and at least one time of day.")
    else:
        ok, med_errors = validate_medication_inputs(
            med_patient, med_name, med_dosage, med_times, med_start, med_end
        )
        if not ok:
            for msg in med_errors:
                st.error(msg)
        else:
            try:
                conn = sqlite3.connect("healthcare.db")
                c = conn.cursor()
                c.execute("""
                    INSERT INTO medications
                    (patient_name, medication_name, dosage, frequency, time_of_day, start_date, end_date)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """, (
                    med_patient,
                    med_name,
                    med_dosage,
                    med_freq,
                    ", ".join(med_times),
                    med_start.strftime("%Y-%m-%d"),
                    med_end.strftime("%Y-%m-%d"),
                ))
                conn.commit()
                new_id = c.lastrowid
                conn.close()
                st.success(f"✅ Medication added to schedule! (ID: {new_id})")
            except Exception as e:
                st.error(f"Could not save medication schedule: {e}")

st.subheader("📋 Today’s Medication Schedule")

if not med_patient:
    st.info("Enter patient name above to see today's schedule.")
else:
    conn = sqlite3.connect("healthcare.db")
    c = conn.cursor()
    c.execute("""
        SELECT id, patient_name, medication_name, dosage,
               frequency, time_of_day, start_date, end_date
        FROM medications
        WHERE LOWER(patient_name) = LOWER(?)
    """, (med_patient.strip(),))
    meds = c.fetchall()

    if not meds:
        st.info("No medications found for this patient.")
    else:
        today_str = date.today().strftime("%Y-%m-%d")

        for med in meds:
            med_id = med[0]
            med_name = med[2]
            dosage = med[3]
            times_str = med[5]
            times_list = [t.strip() for t in times_str.split(",") if t.strip()]

            # logs for this medication
            c.execute("""
                SELECT id, medication_id, log_date, time_of_day, status
                FROM medication_logs
                WHERE medication_id = ?
            """, (med_id,))
            med_logs = c.fetchall()

            expected, taken, adherence = compute_adherence(med, med_logs)

            st.markdown(f"### 💊 {med_name} ({dosage})")
            st.write(f"Time(s) of day: {', '.join(times_list)}")
            st.write(f"Period: {med[6]} → {med[7]}")
            st.write(
                f"Adherence so far: **{taken}/{expected} doses** "
                f"({adherence:.1f}%)"
            )

            # For each time-of-day today, check if already taken
            for t in times_list:
                already = any(
                    log[2] == today_str and log[3] == t and log[4] == "taken"
                    for log in med_logs
                )

                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(
                        f"- {t} dose ({today_str}) — "
                        f"{'✅ Taken' if already else '⏳ Pending'}"
                    )
                with col2:
                    if not already:
                        if st.button(
                            f"Mark {med_name} {t} as taken",
                            key=f"{med_id}_{t}"
                        ):
                            c.execute("""
                                INSERT INTO medication_logs
                                (medication_id, log_date, time_of_day, status)
                                VALUES (?, ?, ?, 'taken')
                            """, (med_id, today_str, t))
                            conn.commit()
                            st.success(f"Marked {med_name} ({t}) as taken!")

    conn.close()



# =====================================================
# MEDICATION ADHERENCE OVERVIEW (PER PATIENT)
# =====================================================

st.markdown("---")
st.header("📈 Medication Adherence Overview")

overview_patient = st.text_input("Patient name (for adherence overview)")

if st.button("Show Adherence Overview"):
    if not overview_patient.strip():
        st.error("Please enter a patient name.")
    else:
        summaries, overall_rate, total_expected = overall_adherence_for_patient(
            overview_patient
        )
        if not summaries:
            st.info("No medications scheduled for this patient.")
        else:
            total_taken = sum(s["taken"] for s in summaries)
            st.write(f"Overall adherence for **{overview_patient}**:")
            st.progress(min(overall_rate / 100.0, 1.0))
            st.write(
                f"**{overall_rate:.1f}%** doses taken ({total_taken}/{total_expected})"
            )

            # bar chart per medication
            data = {
                "Medication": [s["name"] for s in summaries],
                "Adherence %": [round(s["adherence"], 1) for s in summaries],
                "Taken/Expected": [f"{s['taken']}/{s['expected']}" for s in summaries],
            }
            df = pd.DataFrame(data).set_index("Medication")
            st.bar_chart(df[["Adherence %"]])

            st.write("Details:")
            for s in summaries:
                status = "✅ Good" if s["adherence"] >= 80 else "⚠ Needs attention"
                st.write(
                    f"- **{s['name']} ({s['dosage']})** — "
                    f"{s['adherence']:.1f}% ({s['taken']}/{s['expected']}) · {status}"
                )
# =====================================================
# MEDICATION CALENDAR VIEW
# =====================================================

st.markdown("---")
st.header("🗓 Medication Calendar View")

cal_patient = st.text_input("Patient name (for calendar view)")

if st.button("Show Medication Calendar"):
    if not cal_patient.strip():
        st.error("Please enter patient name.")
    else:
        df_cal = build_medication_calendar_df(cal_patient)
        if df_cal is None:
            st.info("No medications scheduled for this patient.")
        else:
            today = date.today()
            st.write(f"Month view for **{today.strftime('%B %Y')}**")
            st.caption("Legend: ✅ taken · ❌ missed · • upcoming · — no medication")

            st.dataframe(df_cal, use_container_width=True)

# =====================================================
# DATA EXPORT CENTER (MEDICATION + WELLNESS)
# =====================================================

st.markdown("---")
st.header("📤 Data Export Center")

export_patient = st.text_input("Patient name (for data export)")

if st.button("Prepare Export Files"):
    if not export_patient.strip():
        st.error("Please enter a patient name.")
    else:
        conn = sqlite3.connect("healthcare.db")
        c = conn.cursor()

        # Medication history with logs
        c.execute("""
            SELECT m.id, m.patient_name, m.medication_name, m.dosage,
                   m.frequency, m.time_of_day, m.start_date, m.end_date,
                   l.log_date, l.time_of_day, l.status
            FROM medications m
            LEFT JOIN medication_logs l
              ON m.id = l.medication_id
            WHERE LOWER(m.patient_name) = LOWER(?)
        """, (export_patient.strip(),))
        med_rows = c.fetchall()

        # Wellness / fitness logs
        c.execute("""
            SELECT date, steps, sleep_hours, water_intake, mood
            FROM wellness_logs
            WHERE LOWER(patient_name) = LOWER(?)
            ORDER BY date
        """, (export_patient.strip(),))
        wellness_rows = c.fetchall()

        conn.close()

        if not med_rows and not wellness_rows:
            st.info("No medication or wellness data found for this patient.")
        else:
            if med_rows:
                med_df = pd.DataFrame(
                    med_rows,
                    columns=[
                        "medication_id", "patient_name", "medication_name", "dosage",
                        "frequency", "time_of_day", "start_date", "end_date",
                        "log_date", "log_time_of_day", "log_status",
                    ],
                )
                med_csv = med_df.to_csv(index=False).encode("utf-8")
                st.download_button(
                    label="⬇ Download Medication History (CSV)",
                    data=med_csv,
                    file_name=f"{export_patient.replace(' ', '_')}_medication_history.csv",
                    mime="text/csv",
                )

            if wellness_rows:
                well_df = pd.DataFrame(
                    wellness_rows,
                    columns=["date", "steps", "sleep_hours", "water_intake", "mood"],
                )
                well_csv = well_df.to_csv(index=False).encode("utf-8")
                st.download_button(
                    label="⬇ Download Wellness / Fitness Logs (CSV)",
                    data=well_csv,
                    file_name=f"{export_patient.replace(' ', '_')}_wellness_logs.csv",
                    mime="text/csv",
                )



# =====================================================
# WELLNESS GOALS & PROGRESS TRACKING
# =====================================================

st.markdown("---")
st.header("🏃 Wellness Goals & Progress")

wg_patient = st.text_input("Patient name (for wellness tracking)")

c1, c2, c3 = st.columns(3)
with c1:
    steps_goal = st.number_input(
        "Daily steps goal", min_value=0, step=500, value=5000
    )
with c2:
    sleep_goal = st.number_input(
        "Sleep goal (hours)", min_value=0.0, max_value=24.0, value=7.0
    )
with c3:
    water_goal = st.number_input(
        "Water goal (ml)", min_value=0, step=250, value=2000
    )

if st.button("Save Wellness Goals"):
    if not wg_patient.strip():
        st.error("Enter patient name.")
    else:
        upsert_wellness_goal(wg_patient, steps_goal, sleep_goal, water_goal)
        st.success("Wellness goals saved!")

st.subheader("Log Today’s Wellness")

c1, c2, c3, c4 = st.columns(4)
with c1:
    steps_today = st.number_input("Steps today", min_value=0, step=100)
with c2:
    sleep_today = st.number_input("Sleep today (hours)", min_value=0.0, max_value=24.0)
with c3:
    water_today = st.number_input("Water today (ml)", min_value=0, step=250)
with c4:
    mood_today = st.selectbox(
        "Mood", ["Very good", "Good", "Okay", "Low", "Stressed"]
    )

if st.button("Save Today’s Wellness Log"):
    if not wg_patient.strip():
        st.error("Enter patient name.")
    else:
        log_wellness_entry(
            wg_patient, steps_today, sleep_today, water_today, mood_today
        )
        st.success("Wellness log saved for today!")

st.subheader("📈 Wellness Progress (last few entries)")

if st.button("Show Wellness Progress"):
    if not wg_patient.strip():
        st.error("Enter patient name.")
    else:
        conn = sqlite3.connect("healthcare.db")
        c = conn.cursor()
        c.execute(
            """
            SELECT date, steps, sleep_hours, water_intake
            FROM wellness_logs
            WHERE LOWER(patient_name) = LOWER(?)
            ORDER BY date
            """,
            (wg_patient.strip(),),
        )
        rows = c.fetchall()
        conn.close()

        if not rows:
            st.info("No wellness logs found for this patient.")
        else:
            df = pd.DataFrame(rows, columns=["date", "steps", "sleep", "water"])
            df["date"] = pd.to_datetime(df["date"])
            df = df.sort_values("date").set_index("date")

            st.line_chart(df[["steps", "water"]])
            st.line_chart(df[["sleep"]])


# =====================================================
# HEALTH INSIGHTS & RECOMMENDATIONS
# =====================================================

st.markdown("---")
st.header("🧠 Health Insights & Recommendations")

insight_patient = st.text_input("Patient name (for insights)")

if st.button("Generate Insights"):
    if not insight_patient.strip():
        st.error("Please enter patient name.")
    else:
        insights = generate_insights(insight_patient)
        for idx, text in enumerate(insights, start=1):
            st.write(f"{idx}. {text}")


# =====================================================
# RISK TIMELINE GRAPH
# =====================================================

st.markdown("---")
st.header("📉 Risk Timeline (Visits History)")

risk_patient = st.text_input("Patient name (for risk timeline)")

if st.button("Show Risk Timeline"):
    if not risk_patient.strip():
        st.error("Please enter patient name.")
    else:
        df_risk = get_risk_history(risk_patient)
        if df_risk is None:
            st.info("No previous risk records found for this patient.")
        else:
            st.line_chart(df_risk[["RiskScore"]])
            st.caption(
                "1 = Low risk, 2 = Moderate, 3 = High. Each point represents one saved health analysis."
            )

# =====================================================
# DAILY SUMMARY CARD + SMART RECOMMENDATIONS
# =====================================================

st.markdown("---")
st.header("📅 Daily Summary & Smart Suggestions")

summary_patient = st.text_input("Patient name (for daily summary)")

if st.button("Show Today’s Summary"):
    if not summary_patient.strip():
        st.error("Please enter patient name.")
    else:
        summary_text, raw = build_daily_summary(summary_patient)
        st.subheader("Today’s Health Snapshot")
        st.info(summary_text)

        st.subheader("Smart Suggestions for Today")
        recos = generate_smart_recommendations(summary_patient)
        for r in recos:
            st.write(f"- {r}")



# =====================================================
# PATIENT HEALTH DASHBOARD (SCORE + GAUGE + ALERTS)
# =====================================================

st.markdown("---")
st.header("📊 Patient Health Dashboard (Score, Gauge & Trends)")

dash_patient = st.text_input("Patient name (for dashboard)")

if st.button("Show Dashboard"):
    if not dash_patient.strip():
        st.error("Please enter patient name.")
    else:
        # latest vitals
        vitals = get_latest_vitals_for_dashboard(dash_patient)

        # medication adherence
        summaries, adherence_rate, total_expected = overall_adherence_for_patient(
            dash_patient
        )
        if total_expected == 0:
            adherence_rate_display = None
        else:
            adherence_rate_display = adherence_rate

        # wellness stats
        wellness_stats, wellness_df = get_wellness_stats_for_dashboard(dash_patient)

        if not vitals and adherence_rate_display is None and not wellness_stats:
            st.info("No data found for this patient yet.")
        else:
            # ---------- HEALTH SCORE ----------
            risk_component = 0
            risk_label = vitals["risk"] if vitals and vitals.get("risk") else None
            if risk_label == "Low":
                risk_component = 40
            elif risk_label == "Moderate":
                risk_component = 25
            elif risk_label == "High":
                risk_component = 10

            adherence_component = 0
            if adherence_rate_display is not None:
                adherence_component = max(0.0, min(adherence_rate_display, 100.0)) * 0.35

            wellness_component = 0
            avg_goal_completion = None
            if wellness_stats and wellness_stats["avg_goal_completion"] is not None:
                avg_goal_completion = wellness_stats["avg_goal_completion"]
                wellness_component = max(0.0, min(avg_goal_completion, 100.0)) * 0.25

            health_score = int(
                round(risk_component + adherence_component + wellness_component)
            )
            health_score = max(0, min(health_score, 100))

            if health_score >= 85:
                score_label = "Excellent 🌟"
            elif health_score >= 70:
                score_label = "Good 🙂"
            elif health_score >= 50:
                score_label = "Needs attention ⚠️"
            else:
                score_label = "High risk ❗"

            # ---------- SCORE + RISK GAUGE ----------
            st.subheader("Overall Health Score")
            st.metric("Health Score", f"{health_score}/100")
            st.progress(health_score / 100.0)
            st.caption(f"Status: {score_label}")

            st.subheader("Risk Gauge")
            if risk_label:
                risk_map = {"Low": 0.3, "Moderate": 0.65, "High": 0.9}
                risk_val = risk_map.get(risk_label, 0.5)
                st.progress(risk_val)
                st.caption(f"Current clinical risk: **{risk_label}**")
            else:
                st.info("No risk level available yet for this patient.")

            # ---------- SMART ALERTS ----------
            st.subheader("⚠️ Smart Health Alerts")
            alerts = []

            if health_score < 50 or risk_label == "High":
                alerts.append(
                    "Overall risk appears high; consider consulting a doctor soon."
                )

            if adherence_rate_display is not None and adherence_rate_display < 80:
                alerts.append(
                    "Medication adherence is below 80%; missed doses may reduce treatment effectiveness."
                )

            if wellness_stats:
                if wellness_stats["avg_sleep"] < 6:
                    alerts.append(
                        "Average sleep is less than 6 hours; improve sleep schedule for better recovery."
                    )
                if wellness_stats["avg_steps"] < 3000:
                    alerts.append(
                        "Physical activity is low; try to walk more each day if your doctor allows."
                    )
                if wellness_stats["avg_water"] < 1500:
                    alerts.append(
                        "Water intake seems low; aim for better hydration unless restricted."
                    )

            if alerts:
                for a in alerts:
                    st.write("• " + a)
            else:
                st.success("No urgent alerts detected. Keep following your healthy routine!")

            # ---------- VITALS SNAPSHOT ----------
            st.subheader("Latest Vitals Snapshot")
            if vitals:
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.write(f"**BP:** {vitals['bp']}")
                with col2:
                    st.write(f"**HR:** {vitals['heart_rate']} bpm")
                with col3:
                    st.write(f"**Temp:** {vitals['temperature']} °C")
                st.write(f"**Risk level:** {vitals['risk']}")
                st.write(f"**Symptoms:** {vitals['symptoms']}")
            else:
                st.info("No vitals saved yet for this patient.")

            # ---------- MEDICATION ADHERENCE ----------
            st.subheader("Medication Adherence Summary")
            if adherence_rate_display is None:
                st.info("No medication schedule found for this patient.")
            else:
                st.write(f"Overall adherence: **{adherence_rate_display:.1f}%**")
                if summaries:
                    data = {
                        "Medication": [s["name"] for s in summaries],
                        "Adherence %": [round(s["adherence"], 1) for s in summaries],
                    }
                    df_dash = pd.DataFrame(data).set_index("Medication")
                    st.bar_chart(df_dash)

            # ---------- WELLNESS TRENDS ----------
            st.subheader("Wellness Trends (last entries)")
            if wellness_stats and wellness_df is not None:
                st.write(
                    f"Average steps: **{wellness_stats['avg_steps']:.0f}** | "
                    f"Average sleep: **{wellness_stats['avg_sleep']:.1f} h** | "
                    f"Average water: **{wellness_stats['avg_water']:.0f} ml**"
                )
                st.write(
                    f"Wellness streak: **{wellness_stats['streak_days']}** "
                    f"day(s) with logs in a row"
                )

                st.line_chart(
                    wellness_df.set_index("date")[["steps", "water"]]
                )
                st.line_chart(
                    wellness_df.set_index("date")[["sleep"]]
                )
            else:
                st.info("No wellness logs yet for this patient.")

            # ---------- INSIGHTS ----------
            st.subheader("AI Insights")
            insights_for_dash = generate_insights(dash_patient)
            for idx, text in enumerate(insights_for_dash, start=1):
                st.write(f"{idx}. {text}")

            # ---------- EXPORTS ----------
            st.subheader("Export Dashboard Summary")
            dash_docx_buf = create_dashboard_docx(
                dash_patient,
                health_score,
                score_label,
                vitals,
                adherence_rate_display,
                wellness_stats,
                insights_for_dash,
            )

            st.download_button(
                label="⬇ Download Dashboard Summary (DOCX)",
                data=dash_docx_buf,
                file_name=f"{dash_patient.replace(' ', '_')}_dashboard_summary.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )

            weekly_pdf_buf = create_weekly_pdf_report(
                dash_patient,
                vitals,
                adherence_rate_display,
                wellness_stats,
                insights_for_dash,
            )
            st.download_button(
                label="⬇ Download Weekly Progress (PDF)",
                data=weekly_pdf_buf,
                file_name=f"{dash_patient.replace(' ', '_')}_weekly_progress.pdf",
                mime="application/pdf",
            )

# =====================================================
# AI LIFESTYLE COACH (DIET + EXERCISE SUGGESTIONS)
# =====================================================

st.markdown("---")
st.header("🧠 AI Lifestyle Coach")

coach_patient = st.text_input("Patient name (for lifestyle coach)")

if st.button("Generate Lifestyle Plan"):
    if not coach_patient.strip():
        st.error("Please enter patient name.")
    else:
        plan = generate_lifestyle_plan_for_patient(coach_patient)
        st.write(
            "These are general suggestions based on recorded vitals, adherence "
            "and wellness trends. They do not replace a doctor's advice."
        )

        for section, items in plan.items():
            st.subheader(section)
            if not items:
                st.write("- (No specific suggestions yet)")
            else:
                for item in items:
                    st.write(f"- {item}")


# =====================================================
# INDIAN CONTEXT HEALTH TIPS
# =====================================================

st.markdown("---")
st.header("🇮🇳 Indian Lifestyle Tips (Context Aware)")

india_patient = st.text_input("Patient name (for India-specific tips)")

if st.button("Show Indian Context Tips"):
    if not india_patient.strip():
        st.error("Please enter patient name.")
    else:
        tips = generate_indian_context_tips(india_patient)
        for t in tips:
            st.write(f"- {t}")



# =====================================================
# FAMILY HEALTH MONITORING & CAREGIVER NOTIFICATIONS
# =====================================================

st.markdown("---")
st.header("👨‍👩‍👧 Family Monitoring & Caregiver Notification")

patient_for_summary = st.text_input("Patient name (whose health to summarize)")
caregiver_name = st.text_input("Primary caregiver name")
caregiver_list_raw = st.text_input(
    "Other caregivers / family members (comma-separated, e.g., Abc (mother), Efg (father))"
)

if st.button("Generate Caregiver Summary"):
    if not patient_for_summary.strip() or not caregiver_name.strip():
        st.error("Please enter both patient name and primary caregiver name.")
    else:
        # Build caregiver list (for info only, not used for insights)
        caregivers = [n.strip() for n in caregiver_list_raw.split(",") if n.strip()]

        # Get health insights for the ONE patient
        insights = generate_insights(patient_for_summary)

        # Caregiver list text
        if caregivers:
            caregivers_block = "\n".join(f"- {n}" for n in caregivers)
        else:
            caregivers_block = "- (not specified)"

        # Build the message text
        insights_block_lines = [f"{idx}. {text}" for idx, text in enumerate(insights, start=1)]
        insights_block = "\n".join(insights_block_lines)

        summary_text = f"""Dear {caregiver_name},

Here is a summary of {patient_for_summary}'s recent health status.

Caregivers on record:
{caregivers_block}

Health insights:
{insights_block}

These points are automatically generated from recorded vitals,
medication adherence, and wellness logs. For any serious concerns,
please consult a qualified doctor.

— Healthcare Monitoring AI Agent
"""

        st.subheader("Caregiver Message (copy & send)")
        st.text_area("Message:", summary_text, height=350)


# =====================================================
# SIMPLE HEALTH Q&A ASSISTANT
# =====================================================

st.markdown("---")
st.header("💬 Health Q&A Assistant")

chat_patient = st.text_input(
    "Patient name (optional – helps personalise steps/sleep/water replies)",
    key="chat_patient_name",
)
chat_question = st.text_input(
    "Ask a general health question (no emergencies, no prescriptions):",
    key="chat_question",
)

if st.button("Ask Health Assistant"):
    if not chat_question.strip():
        st.error("Please type a question.")
    else:
        answers = answer_health_question(chat_question, chat_patient or None)
        for ans in answers:
            st.write("- " + ans)

